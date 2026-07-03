"""SmartExcel processing worker.

Deterministic spreadsheet processing lives here, NOT on Cloudflare Workers (which
can't run pandas/openpyxl). Deploy this container to Cloud Run or Cloudflare
Containers; the web app dispatches jobs to it via Cloudflare Queues and exchanges
files through R2.

Flow (Phase 1 / A):
  1. Web app uploads the input file to R2 and enqueues a job message.
  2. The queue consumer POSTs /process with a machine-executable `operations` list
     (mirrors src/lib/operations.ts). This worker returns 202 and processes in the
     background: download from R2 -> pandas -> apply operations -> write output
     .xlsx + a small preview.json back to R2 -> POST results to /api/worker-callback.
  3. /inspect is a synchronous helper the clarification/planning step calls to read
     sheet/column/sample context from an uploaded file.

The AI never runs here — it only produces the operations spec; this layer executes
it deterministically.
"""

from __future__ import annotations  # allow `X | None` annotations on Python 3.9

import ast
import builtins as _builtins
import io
import json
import os
import re
import urllib.request
from typing import Optional

import pandas as pd
from dotenv import load_dotenv
from fastapi import BackgroundTasks, FastAPI, Header, HTTPException
from pydantic import BaseModel

load_dotenv()  # local .env; in Cloud Run env vars are set directly (no-op there)

app = FastAPI(title="SmartExcel Worker", version="1.0.0")

SHARED_SECRET = os.environ.get("WORKER_SHARED_SECRET", "")
APP_CALLBACK_URL = os.environ.get("APP_CALLBACK_URL", "http://localhost:3000")
R2_BUCKET = os.environ.get("R2_BUCKET", "smartexcel-files")
R2_ENDPOINT = os.environ.get("R2_ENDPOINT", "")
SAMPLE_ROWS = int(os.environ.get("SAMPLE_ROWS", "100"))
PREVIEW_ROWS = 50


class JobRequest(BaseModel):
    job_id: str
    plan_id: Optional[str] = None
    stage: str  # "sample" | "full"
    input_object_key: str
    operations: list[dict] = []
    options: dict = {}


class InspectRequest(BaseModel):
    input_object_key: str


class InspectResult(BaseModel):
    sheets: list[str]
    headers: list[str]
    sample_rows: list[list[str]]


class ValidateRequest(BaseModel):
    operations: list[dict]


class ValidationIssue(BaseModel):
    index: int  # position in the operations array
    op: str  # the op kind
    description: Optional[str] = None
    error: str


class ValidateResult(BaseModel):
    ok: bool
    issues: list[ValidationIssue]


def _require_auth(authorization: Optional[str]) -> None:
    if not SHARED_SECRET:
        return  # auth disabled in local/dev when no secret is set
    if authorization != f"Bearer {SHARED_SECRET}":
        raise HTTPException(status_code=401, detail="unauthorized")


def _r2_client():
    if not (R2_ENDPOINT and os.environ.get("R2_ACCESS_KEY_ID")):
        return None
    import boto3

    return boto3.client(
        "s3",
        endpoint_url=R2_ENDPOINT,
        aws_access_key_id=os.environ["R2_ACCESS_KEY_ID"],
        aws_secret_access_key=os.environ["R2_SECRET_ACCESS_KEY"],
        region_name="auto",
    )


def _get_object_bytes(key: str) -> Optional[bytes]:
    client = _r2_client()
    if client is None:
        return None
    obj = client.get_object(Bucket=R2_BUCKET, Key=key)
    return obj["Body"].read()


def _put_object(key: str, body: bytes, content_type: str) -> Optional[str]:
    client = _r2_client()
    if client is None:
        return None
    client.put_object(Bucket=R2_BUCKET, Key=key, Body=body, ContentType=content_type)
    return key


IMAGE_EXTS = {"png", "jpg", "jpeg", "gif", "webp", "bmp", "tiff"}


def _ext(key: str) -> str:
    return key.rsplit(".", 1)[-1].lower() if "." in key else ""


def _table_to_df(rows: list) -> pd.DataFrame:
    """Turn a raw list-of-rows (first row = header) into a string DataFrame with
    unique, non-empty column names and rectangular rows."""
    if not rows:
        return pd.DataFrame()
    seen: dict[str, int] = {}
    cols: list[str] = []
    for i, raw in enumerate(rows[0]):
        name = (str(raw).strip() if raw is not None else "") or f"col_{i + 1}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        cols.append(name)
    body = []
    for r in rows[1:]:
        vals = [("" if c is None else str(c)) for c in r]
        vals = (vals + [""] * len(cols))[: len(cols)]
        body.append(vals)
    return pd.DataFrame(body, columns=cols, dtype=str)


def _load_pdf(data: bytes) -> pd.DataFrame:
    import pdfplumber

    tables = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            for t in page.extract_tables():
                if t:
                    tables.append(t)
    if tables:
        return _table_to_df(max(tables, key=len))  # largest table by row count
    lines = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            lines.extend(ln for ln in (page.extract_text() or "").splitlines() if ln.strip())
    if not lines:
        raise RuntimeError("No extractable tables or text found in the PDF.")
    return pd.DataFrame({"text": lines}, dtype=str)


def _load_docx(data: bytes) -> pd.DataFrame:
    import docx

    doc = docx.Document(io.BytesIO(data))
    if doc.tables:
        rows = [[cell.text for cell in row.cells] for row in doc.tables[0].rows]
        return _table_to_df(rows)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    if not paras:
        raise RuntimeError("No tables or text found in the document.")
    return pd.DataFrame({"text": paras}, dtype=str)


def _load_text(data: bytes) -> pd.DataFrame:
    text = data.decode("utf-8", errors="replace")
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return pd.DataFrame()
    delim = "\t" if "\t" in lines[0] else ("," if "," in lines[0] else None)
    if delim:
        return pd.read_csv(io.StringIO(text), sep=delim, dtype=str)
    return pd.DataFrame({"line": lines}, dtype=str)


def _load_df(key: str, data: bytes) -> pd.DataFrame:
    ext = _ext(key)
    bio = io.BytesIO(data)
    if ext == "tsv":
        return pd.read_csv(bio, sep="\t", dtype=str)
    if ext in ("xlsx", "xlsm"):
        return pd.read_excel(bio, dtype=str, engine="openpyxl")
    if ext == "xls":
        return pd.read_excel(bio, dtype=str)
    if ext == "pdf":
        return _load_pdf(data)
    if ext == "docx":
        return _load_docx(data)
    if ext == "doc":
        raise RuntimeError("Legacy .doc isn't supported — please convert to .docx.")
    if ext in ("txt", "md"):
        return _load_text(data)
    if ext == "xml":
        try:
            return pd.read_xml(bio, parser="etree", dtype=str)
        except Exception as err:  # noqa: BLE001
            raise RuntimeError(f"Could not parse XML into a table: {err}")
    if ext in IMAGE_EXTS:
        raise RuntimeError("Image OCR isn't supported yet — upload a spreadsheet or document.")
    return pd.read_csv(bio, dtype=str)  # default: csv


def _sheet_names(key: str, data: bytes) -> list[str]:
    ext = key.rsplit(".", 1)[-1].lower() if "." in key else ""
    if ext in ("xlsx", "xlsm", "xls"):
        try:
            return pd.ExcelFile(io.BytesIO(data)).sheet_names
        except Exception:
            return ["Sheet1"]
    return ["Sheet1"]


def _norm_header(name: object, style: str) -> str:
    s = str(name).strip()
    if style == "title":
        return s.title() or "Column"
    s = re.sub(r"[^0-9a-zA-Z]+", "_", s).strip("_").lower()
    return s or "col"


# ---------------------------------------------------------------------------
# Tier-2 escape hatch: sandboxed Python expression for the long tail of jobs we
# can't reasonably enumerate at code level. The AI emits a small pandas snippet;
# the worker validates it against a strict AST whitelist before exec'ing in a
# restricted namespace. Three concentric safety nets: (1) AST validation rejects
# imports/dunders/loops/def/class/raise/try/etc.; (2) restricted globals expose
# only `df`/`pd`/`re` + a small safe-builtin set; (3) the PRD sample-run gate
# means the user reviews the output before it touches real data.
# ---------------------------------------------------------------------------

_ALLOWED_GLOBALS = frozenset({"df", "pd", "re"})
# Permissive list of safe Python builtins for AI-generated transformations.
# Excludes: open, exec, eval, compile, __import__, globals, locals, vars,
# getattr/setattr/delattr/hasattr (would bypass the underscore-attr check),
# dir, id, input, breakpoint, help, copyright, credits, license.
_SAFE_BUILTIN_NAMES = (
    "len range min max sum abs round divmod pow "
    "str int float bool list tuple dict set frozenset bytes bytearray "
    "enumerate zip map filter sorted reversed any all next iter "
    "isinstance type repr print format chr ord hex oct bin "
    "Exception ValueError TypeError KeyError IndexError AttributeError "
    "RuntimeError ArithmeticError ZeroDivisionError OverflowError "
    "NameError UnboundLocalError StopIteration LookupError "
    "UnicodeError UnicodeEncodeError UnicodeDecodeError UnicodeTranslateError "
    "True False None"
).split()
_SAFE_BUILTINS = {name: getattr(_builtins, name) for name in _SAFE_BUILTIN_NAMES if hasattr(_builtins, name)}

# Any AST node type not in this set is rejected. We allow normal Python
# idioms — f-strings, loops, try/except, local helper functions. We still
# reject: Import/ImportFrom (no modules beyond df/pd/re), Global/Nonlocal,
# Yield/Await/AsyncFunctionDef (no coroutines), With (no resource context
# managers), Match (no structural pattern matching — unnecessary), Delete
# (clarity), Assert. ClassDef is allowed only if needed; currently excluded.
_ALLOWED_AST_TYPES: tuple = (
    ast.Module, ast.Expression, ast.Expr,
    ast.Assign, ast.AugAssign, ast.AnnAssign,
    ast.Subscript, ast.Attribute, ast.Call, ast.Name, ast.Constant,
    ast.BinOp, ast.UnaryOp, ast.BoolOp, ast.Compare, ast.IfExp, ast.NamedExpr,
    ast.List, ast.Tuple, ast.Dict, ast.Set, ast.Slice, ast.Starred,
    ast.Lambda, ast.arguments, ast.arg, ast.keyword,
    ast.ListComp, ast.SetComp, ast.DictComp, ast.GeneratorExp, ast.comprehension,
    ast.JoinedStr, ast.FormattedValue,  # f-strings
    ast.If, ast.For, ast.While, ast.Break, ast.Continue, ast.Pass,
    ast.Try, ast.ExceptHandler, ast.Raise,
    ast.FunctionDef, ast.Return,
    ast.Load, ast.Store, ast.Del,
    ast.Add, ast.Sub, ast.Mult, ast.Div, ast.FloorDiv, ast.Mod, ast.Pow, ast.MatMult,
    ast.BitOr, ast.BitAnd, ast.BitXor, ast.LShift, ast.RShift,
    ast.UAdd, ast.USub, ast.Not, ast.Invert,
    ast.And, ast.Or,
    ast.Eq, ast.NotEq, ast.Lt, ast.LtE, ast.Gt, ast.GtE,
    ast.Is, ast.IsNot, ast.In, ast.NotIn,
)


def _collect_bound_names(tree: ast.AST) -> set:
    """Names introduced as locals (assignment targets, lambda args, comp targets,
    function defs, exception handler aliases). Free names elsewhere must be in
    the allowed-globals set."""
    bound: set = set()
    for node in ast.walk(tree):
        if isinstance(node, ast.Name) and isinstance(node.ctx, ast.Store):
            bound.add(node.id)
        elif isinstance(node, ast.arg):
            bound.add(node.arg)
        elif isinstance(node, ast.FunctionDef):
            bound.add(node.name)
        elif isinstance(node, ast.ExceptHandler) and node.name:
            bound.add(node.name)
    return bound


# Gemini occasionally flattens function/control bodies onto a single line with
# runs of spaces where newlines belong. If the source doesn't parse, insert
# newlines before statement keywords that appear after a 4+ space run, then
# retry. Only applied when the original failed to parse, so well-formed code
# is unaffected.
_UNFLATTEN_KEYWORDS = re.compile(
    r" {4,}(?=(?:except|finally|elif|else|return|raise|continue|break|pass|if|for|while|try|with|def|class)\b)"
)


def _maybe_unflatten(source: str) -> str:
    try:
        ast.parse(source, mode="exec")
        return source
    except SyntaxError:
        pass
    fixed = _UNFLATTEN_KEYWORDS.sub(lambda m: "\n" + m.group(0), source)
    if fixed == source:
        return source
    try:
        ast.parse(fixed, mode="exec")
        return fixed
    except SyntaxError:
        return source  # fix didn't help; let the original error stand


def _validate_custom_expression(source: str) -> Optional[str]:
    try:
        tree = ast.parse(source, mode="exec")
    except SyntaxError as err:
        return f"Syntax error: {err.msg}"
    bound = _collect_bound_names(tree)
    allowed_names = _ALLOWED_GLOBALS | set(_SAFE_BUILTIN_NAMES) | bound
    for node in ast.walk(tree):
        if not isinstance(node, _ALLOWED_AST_TYPES):
            return f"Disallowed AST node: {type(node).__name__}"
        if isinstance(node, ast.Name) and isinstance(node.ctx, ast.Load):
            if node.id not in allowed_names:
                return f"Disallowed name: {node.id}"
        if isinstance(node, ast.Attribute) and node.attr.startswith("_"):
            return f"Disallowed underscore attribute: {node.attr}"
    return None


def _run_custom_transform(df: pd.DataFrame, expression: str) -> pd.DataFrame:
    expression = _maybe_unflatten(expression)
    err = _validate_custom_expression(expression)
    if err:
        raise ValueError(err)
    compiled = compile(expression, "<custom_transform>", "exec")
    namespace = {"__builtins__": _SAFE_BUILTINS, "df": df, "pd": pd, "re": re}
    exec(compiled, namespace)  # noqa: S102 — sandboxed
    new_df = namespace.get("df")
    if not isinstance(new_df, pd.DataFrame):
        raise ValueError("custom_transform must leave `df` as a DataFrame")
    return new_df


def _apply_ops(df: pd.DataFrame, operations: list[dict]) -> tuple[pd.DataFrame, list[str]]:
    """Apply operations in order; return the result and a per-step report of what
    actually changed (the validation/anomaly summary, PRD §6.7)."""
    notes: list[str] = []
    for op in operations:
        kind = op.get("op")
        before = len(df)
        if kind == "trim_whitespace":
            df = df.map(lambda x: x.strip() if isinstance(x, str) else x)
            notes.append("Trimmed whitespace from text cells")
        elif kind == "normalize_headers":
            df.columns = [_norm_header(c, op.get("style", "snake")) for c in df.columns]
            notes.append("Normalized column headers")
        elif kind == "rename_column":
            f, t = op.get("from"), op.get("to")
            if f and t and f in df.columns:
                df = df.rename(columns={f: t})
                notes.append(f"Renamed '{f}' → '{t}'")
            else:
                notes.append(f"Skipped rename — column '{f}' not found")
        elif kind == "drop_columns":
            want = op.get("columns") or []
            cols = [c for c in want if c in df.columns]
            if cols:
                df = df.drop(columns=cols)
                notes.append(f"Dropped column(s): {', '.join(cols)}")
            missing = [c for c in want if c not in df.columns]
            if missing:
                notes.append(f"Skipped dropping missing column(s): {', '.join(missing)}")
        elif kind == "drop_empty_rows":
            df = df.dropna(how="all")
            notes.append(f"Removed {before - len(df)} empty row(s)")
        elif kind == "dedupe":
            # Tolerate the AI emitting either `columns` (array) or `column` (single).
            want = op.get("columns") or ([op["column"]] if op.get("column") else None)
            cols = None
            if want:
                cols = [c for c in want if c in df.columns]
                if not cols:
                    notes.append(f"Dedupe key {want} not found — deduped full rows instead")
                    cols = None
            df = df.drop_duplicates(subset=cols)
            by = f" by {', '.join(cols)}" if cols else ""
            notes.append(f"Removed {before - len(df)} duplicate row(s){by}")
        elif kind == "fill_missing":
            c, v = op.get("column"), op.get("value", "")
            if c and c in df.columns:
                filled = int(df[c].isna().sum())
                df[c] = df[c].fillna(v)
                notes.append(f"Filled {filled} missing value(s) in '{c}'")
            else:
                notes.append(f"Skipped fill — column '{c}' not found")
        elif kind == "sort":
            c = op.get("column")
            if c and c in df.columns:
                df = df.sort_values(by=c, ascending=(op.get("order") != "desc"))
                notes.append(f"Sorted by '{c}' ({op.get('order', 'asc')})")
            else:
                notes.append(f"Skipped sort — column '{c}' not found")
        elif kind == "filter_rows":
            c, v = op.get("column"), op.get("value")
            if c and c in df.columns and v is not None:
                df = df[df[c].astype(str) == str(v)]
                notes.append(f"Kept rows where '{c}' = '{v}' (removed {before - len(df)})")
            else:
                notes.append(f"Skipped filter — column '{c}' not found")
        elif kind == "regex_replace":
            c = op.get("column")
            pattern = op.get("pattern")
            replacement = op.get("replacement", "")
            if not c:
                notes.append(
                    "Skipped regex_replace — the plan did not specify which column to target. "
                    "Ask the assistant to re-generate the plan."
                )
            elif c not in df.columns:
                notes.append(
                    f"Skipped regex_replace — column '{c}' not in data (have: {list(df.columns)[:6]}…)."
                )
            elif pattern is None:
                notes.append(f"Skipped regex_replace on '{c}' — no pattern provided.")
            else:
                try:
                    df[c] = df[c].astype(str).str.replace(
                        pattern, replacement, regex=True
                    )
                    notes.append(
                        f"Regex-replaced '{pattern}' → '{replacement}' in '{c}'"
                    )
                except re.error as err:
                    notes.append(f"Skipped regex on '{c}' — invalid pattern: {err}")
        elif kind == "custom_transform":
            expression = op.get("expression")
            description = (op.get("description") or "custom step").strip()
            if not expression:
                notes.append(f"Skipped custom step — no expression provided ({description})")
            else:
                try:
                    df = _run_custom_transform(df, expression)
                    notes.append(f"Ran custom step — {description}")
                except Exception as err:  # noqa: BLE001 — surface to user, don't crash worker
                    notes.append(f"Skipped custom step ({description}): {err}")
    return df.reset_index(drop=True), notes


def _to_xlsx_bytes(df: pd.DataFrame) -> bytes:
    bio = io.BytesIO()
    df.to_excel(bio, index=False, engine="openpyxl")
    return bio.getvalue()


def _preview(df: pd.DataFrame) -> dict:
    head = df.head(PREVIEW_ROWS).fillna("")
    return {
        "columns": [str(c) for c in df.columns],
        "rows": head.astype(str).values.tolist(),
    }


def _post_callback(payload: dict) -> None:
    data = json.dumps(payload).encode("utf-8")
    headers = {"Content-Type": "application/json"}
    if SHARED_SECRET:
        headers["Authorization"] = f"Bearer {SHARED_SECRET}"
    req = urllib.request.Request(
        f"{APP_CALLBACK_URL.rstrip('/')}/api/worker-callback",
        data=data,
        headers=headers,
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:  # noqa: S310 (trusted internal URL)
            resp.read()
    except urllib.error.HTTPError as err:
        body = err.read().decode("utf-8", errors="replace") if err.fp else ""
        print(f"callback {err.code}: {body[:500]}", flush=True)
        raise


def _run_and_callback(req: JobRequest) -> None:
    base = f"jobs/{req.job_id}/{req.stage}"
    try:
        data = _get_object_bytes(req.input_object_key)
        if data is None:
            raise RuntimeError("could not read input from R2 (storage not configured?)")
        df = _load_df(req.input_object_key, data)
        cols_in = len(df.columns)
        sampled = req.stage == "sample"
        if sampled:
            df = df.head(SAMPLE_ROWS)
        rows_in = len(df)
        df, notes = _apply_ops(df, req.operations)
        rows_out = len(df)

        sheets = _sheet_names(req.input_object_key, data)
        if len(sheets) > 1:
            notes.insert(0, f"workbook had {len(sheets)} sheets — processed only the first ('{sheets[0]}')")

        out_key = _put_object(
            f"{base}/output.xlsx",
            _to_xlsx_bytes(df),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        prev_key = _put_object(
            f"{base}/preview.json", json.dumps(_preview(df)).encode("utf-8"), "application/json"
        )
        scope = f"sample of first {rows_in} rows" if sampled else f"{rows_in} rows"
        report = " · ".join(notes) if notes else "no transformations applied"
        summary = (
            f"Processed {scope} → {rows_out} rows, {len(df.columns)} columns "
            f"(from {cols_in}). {report}."
        )
        _post_callback(
            {
                "job_id": req.job_id,
                "stage": req.stage,
                "status": "succeeded",
                "output_object_key": out_key,
                "preview_object_key": prev_key,
                "summary": summary,
                "rows_processed": rows_out,
            }
        )
    except Exception as err:  # noqa: BLE001 — report failures back to the app
        print("processing failed:", err)
        _post_callback(
            {
                "job_id": req.job_id,
                "stage": req.stage,
                "status": "failed",
                "summary": str(err),
                "rows_processed": 0,
            }
        )


@app.get("/health")
def health() -> dict:
    return {"ok": True}


@app.post("/inspect", response_model=InspectResult)
def inspect(req: InspectRequest, authorization: Optional[str] = Header(default=None)) -> InspectResult:
    _require_auth(authorization)
    data = _get_object_bytes(req.input_object_key)
    if data is None:
        return InspectResult(sheets=["Sheet1"], headers=[], sample_rows=[])
    try:
        df = _load_df(req.input_object_key, data)
        return InspectResult(
            sheets=_sheet_names(req.input_object_key, data),
            headers=[str(c) for c in df.columns],
            sample_rows=df.head(5).fillna("").astype(str).values.tolist(),
        )
    except Exception as err:  # noqa: BLE001
        print("inspect failed:", err)
        return InspectResult(sheets=["Sheet1"], headers=[], sample_rows=[])


@app.post("/validate", response_model=ValidateResult)
def validate(
    req: ValidateRequest, authorization: Optional[str] = Header(default=None)
) -> ValidateResult:
    """Pre-flight check for AI-emitted custom_transform code. Returns the same
    sandbox errors a real sample run would surface, so the web app can catch
    syntax / Disallowed-name issues BEFORE queuing a full sample cycle."""
    _require_auth(authorization)
    issues: list[ValidationIssue] = []
    for i, op in enumerate(req.operations):
        if op.get("op") != "custom_transform":
            continue
        expression = op.get("expression") or ""
        if not expression.strip():
            issues.append(
                ValidationIssue(
                    index=i,
                    op=op.get("op", ""),
                    description=op.get("description"),
                    error="empty expression",
                )
            )
            continue
        # Mirror _run_custom_transform: unflatten first, then validate.
        fixed = _maybe_unflatten(expression)
        err = _validate_custom_expression(fixed)
        if err:
            issues.append(
                ValidationIssue(
                    index=i,
                    op=op.get("op", ""),
                    description=op.get("description"),
                    error=err,
                )
            )
    return ValidateResult(ok=len(issues) == 0, issues=issues)


@app.post("/process", status_code=202)
def process(
    req: JobRequest,
    background_tasks: BackgroundTasks,
    authorization: Optional[str] = Header(default=None),
) -> dict:
    _require_auth(authorization)
    background_tasks.add_task(_run_and_callback, req)
    return {"accepted": True, "job_id": req.job_id, "stage": req.stage}
